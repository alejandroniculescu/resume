from __future__ import annotations

import argparse
import textwrap
from datetime import date
from pathlib import Path

import yaml
from docx import Document
from docx.enum.section import WD_SECTION
from docx.shared import Inches, Pt


DEFAULT_SECTION_ORDER = [
    "contracts",
    "private_equity",
    "advisory",
    "work",
    "care",
    "education",
    "volunteer",
    "publications",
]

SECTION_TITLES = {
    "contracts": "CONTRACTS",
    "private_equity": "PRIVATE EQUITY",
    "advisory": "ADVISORY",
    "work": "WORK EXPERIENCE",
    "care": "CARE",
    "education": "EDUCATION",
    "volunteer": "VOLUNTEER",
    "publications": "PUBLICATIONS",
}


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(
        description="Generate ATS-friendly plain text and DOCX resumes from a YAML source file."
    )
    parser.add_argument(
        "input_path",
        nargs="?",
        default="work.yml",
        help="Path to the source YAML resume file.",
    )
    parser.add_argument(
        "--output-dir",
        help="Directory where rendered resume files will be written. Defaults to ./resumes beside the input YAML.",
    )
    parser.add_argument(
        "--base-name",
        default="resume_ats",
        help="Base filename for generated outputs, without extension.",
    )
    parser.add_argument(
        "--sections",
        help=(
            "Comma-separated top-level sections to render in order. "
            "Example: contracts,private_equity,advisory,work,education"
        ),
    )
    parser.add_argument(
        "--has-publications",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Include the publications section. Use --no-has-publications to omit it.",
    )
    parser.add_argument(
        "--has-volunteer",
        action=argparse.BooleanOptionalAction,
        default=True,
        help="Include the volunteer section. Use --no-has-volunteer to omit it.",
    )
    return parser.parse_args()


def load_resume_data(input_path: Path) -> dict:
    with input_path.open("r", encoding="utf-8") as handle:
        data = yaml.safe_load(handle)
    if not isinstance(data, dict):
        raise ValueError("Top-level YAML content must be a mapping.")
    return data


def normalize_text(value: str | None) -> str:
    if not value:
        return ""
    return " ".join(str(value).split())


def title_case_key(key: str) -> str:
    return key.replace("_", " ").upper()


def sanitize_filename_part(value: str) -> str:
    sanitized = "".join(character if character.isalnum() else "_" for character in value.strip())
    collapsed = "_".join(part for part in sanitized.split("_") if part)
    return collapsed or "resume"


def first_name_from_resume(resume_data: dict) -> str:
    about_me = resume_data.get("about_me", {})
    name = normalize_text(about_me.get("name"))
    if not name:
        return "Alexander"
    return name.split()[0]


def resolve_section_order(resume_data: dict, requested_sections: str | None) -> list[str]:
    if requested_sections:
        section_order = [section.strip() for section in requested_sections.split(",") if section.strip()]
    else:
        available_sections = [key for key in resume_data.keys() if key != "about_me"]
        preferred = [section for section in DEFAULT_SECTION_ORDER if section in available_sections]
        remainder = [section for section in available_sections if section not in preferred]
        section_order = preferred + remainder

    missing_sections = [section for section in section_order if section not in resume_data]
    if missing_sections:
        missing = ", ".join(missing_sections)
        raise ValueError(f"Requested sections not found in YAML: {missing}")
    return section_order


def filter_section_order(
    section_order: list[str],
    include_publications: bool,
    include_volunteer: bool,
) -> list[str]:
    filtered_sections: list[str] = []
    for section in section_order:
        if section == "publications" and not include_publications:
            continue
        if section == "volunteer" and not include_volunteer:
            continue
        filtered_sections.append(section)
    return filtered_sections


def format_date_range(item: dict) -> str:
    if item.get("period"):
        return normalize_text(item["period"])

    start_date = item.get("start_date")
    end_date = item.get("end_date")
    if start_date and end_date:
        return f"{start_date} to {end_date}"
    if start_date:
        return f"{start_date} to present"
    if end_date:
        return f"through {end_date}"
    return ""


def wrap_text(text: str, width: int = 96, initial_indent: str = "", subsequent_indent: str = "") -> list[str]:
    normalized = normalize_text(text)
    if not normalized:
        return []
    return textwrap.wrap(
        normalized,
        width=width,
        initial_indent=initial_indent,
        subsequent_indent=subsequent_indent,
        break_long_words=False,
        break_on_hyphens=False,
    )


def append_wrapped(lines: list[str], text: str, indent: str = "", width: int = 96) -> None:
    lines.extend(wrap_text(text, width=width, initial_indent=indent, subsequent_indent=indent))


def role_heading(role: dict) -> str:
    title = normalize_text(role.get("title"))
    role_name = normalize_text(role.get("role"))
    company = normalize_text(role.get("company"))
    organization = normalize_text(role.get("organization"))
    school = normalize_text(role.get("school"))
    publication = normalize_text(role.get("publication"))

    primary = title or role_name or company or organization or school or publication
    secondary = company or organization or school or publication
    if primary and secondary and primary != secondary:
        return f"{primary} | {secondary}"
    return primary or secondary


def render_role_text(lines: list[str], role: dict, include_projects: bool = True) -> None:
    heading = role_heading(role)
    if heading:
        lines.append(heading)

    date_range = format_date_range(role)
    if date_range:
        lines.append(date_range)

    if role.get("organizations"):
        organizations = "; ".join(normalize_text(item) for item in role["organizations"] if normalize_text(item))
        append_wrapped(lines, f"Organizations: {organizations}")

    if role.get("degree"):
        append_wrapped(lines, f"Degree: {role['degree']}")

    if role.get("courses"):
        append_wrapped(lines, f"Courses: {role['courses']}")

    if role.get("publication"):
        append_wrapped(lines, f"Publication: {role['publication']}")

    if role.get("date"):
        append_wrapped(lines, f"Date: {role['date']}")

    if role.get("link"):
        append_wrapped(lines, f"Link: {role['link']}")

    if role.get("description"):
        append_wrapped(lines, role["description"])

    if include_projects and role.get("projects"):
        lines.append("Selected Projects:")
        for project in role["projects"]:
            project_heading_parts = [
                normalize_text(project.get("name")),
                normalize_text(project.get("role")),
            ]
            project_heading = " | ".join(part for part in project_heading_parts if part)
            if project_heading:
                append_wrapped(lines, project_heading, indent="- ", width=94)

            project_date_range = format_date_range(project)
            if project_date_range:
                append_wrapped(lines, project_date_range, indent="  ", width=94)

            if project.get("description"):
                append_wrapped(lines, project["description"], indent="  ", width=94)


def render_text_resume(resume_data: dict, section_order: list[str]) -> str:
    about_me = resume_data.get("about_me", {})
    lines: list[str] = []

    name = normalize_text(about_me.get("name"))
    if name:
        lines.append(name.upper())

    contact_fields = [
        normalize_text(about_me.get("address")),
        normalize_text(about_me.get("email")),
        normalize_text(about_me.get("phone")),
        normalize_text(about_me.get("linkedin")),
        normalize_text(about_me.get("github")),
    ]
    contact_line = " | ".join(field for field in contact_fields if field)
    if contact_line:
        lines.append(contact_line)

    summary = normalize_text(about_me.get("summary"))
    if summary:
        lines.append("")
        lines.append("SUMMARY")
        append_wrapped(lines, summary)

    for section_key in section_order:
        section_value = resume_data.get(section_key)
        if not section_value:
            continue

        lines.append("")
        lines.append(SECTION_TITLES.get(section_key, title_case_key(section_key)))

        if isinstance(section_value, dict) and section_value.get("roles"):
            for index, role in enumerate(section_value["roles"]):
                if index > 0:
                    lines.append("")
                render_role_text(lines, role)
        elif isinstance(section_value, list):
            for index, item in enumerate(section_value):
                if index > 0:
                    lines.append("")
                render_role_text(lines, item)
        else:
            append_wrapped(lines, str(section_value))

    return "\n".join(lines).strip() + "\n"


def set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.start_type = WD_SECTION.CONTINUOUS
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.65)
    section.right_margin = Inches(0.65)

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(10.5)

    for style_name in ["Heading 1", "Heading 2"]:
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.bold = True


def add_paragraph(document: Document, text: str = "", style: str | None = None, bold: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(normalize_text(text))
    run.bold = bold


def render_docx_role(document: Document, role: dict) -> None:
    heading = role_heading(role)
    if heading:
        add_paragraph(document, heading, bold=True)

    date_range = format_date_range(role)
    if date_range:
        add_paragraph(document, date_range)

    if role.get("organizations"):
        organizations = "; ".join(normalize_text(item) for item in role["organizations"] if normalize_text(item))
        add_paragraph(document, f"Organizations: {organizations}")

    if role.get("degree"):
        add_paragraph(document, f"Degree: {role['degree']}")

    if role.get("courses"):
        add_paragraph(document, f"Courses: {role['courses']}")

    if role.get("publication"):
        add_paragraph(document, f"Publication: {role['publication']}")

    if role.get("date"):
        add_paragraph(document, f"Date: {role['date']}")

    if role.get("link"):
        add_paragraph(document, f"Link: {role['link']}")

    if role.get("description"):
        add_paragraph(document, role["description"])

    if role.get("projects"):
        add_paragraph(document, "Selected Projects:", bold=True)
        for project in role["projects"]:
            parts = [normalize_text(project.get("name")), normalize_text(project.get("role"))]
            project_heading = " | ".join(part for part in parts if part)
            if project_heading:
                add_paragraph(document, project_heading, style="List Bullet")
            if project.get("period"):
                add_paragraph(document, project["period"])
            elif project.get("start_date") or project.get("end_date"):
                add_paragraph(document, format_date_range(project))
            if project.get("description"):
                add_paragraph(document, project["description"])


def render_docx_resume(resume_data: dict, section_order: list[str], output_path: Path) -> None:
    document = Document()
    set_document_defaults(document)

    about_me = resume_data.get("about_me", {})
    name = normalize_text(about_me.get("name"))
    if name:
        add_paragraph(document, name, style="Heading 1")

    contact_fields = [
        normalize_text(about_me.get("address")),
        normalize_text(about_me.get("email")),
        normalize_text(about_me.get("phone")),
        normalize_text(about_me.get("linkedin")),
        normalize_text(about_me.get("github")),
    ]
    contact_line = " | ".join(field for field in contact_fields if field)
    if contact_line:
        add_paragraph(document, contact_line)

    summary = normalize_text(about_me.get("summary"))
    if summary:
        add_paragraph(document, "SUMMARY", style="Heading 2")
        add_paragraph(document, summary)

    for section_key in section_order:
        section_value = resume_data.get(section_key)
        if not section_value:
            continue

        add_paragraph(document, SECTION_TITLES.get(section_key, title_case_key(section_key)), style="Heading 2")

        if isinstance(section_value, dict) and section_value.get("roles"):
            for role in section_value["roles"]:
                render_docx_role(document, role)
        elif isinstance(section_value, list):
            for item in section_value:
                render_docx_role(document, item)
        else:
            add_paragraph(document, str(section_value))

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    args = parse_args()
    input_path = Path(args.input_path)
    output_dir = Path(args.output_dir) if args.output_dir else input_path.resolve().parent / "resumes"
    output_dir.mkdir(parents=True, exist_ok=True)

    resume_data = load_resume_data(input_path)
    section_order = resolve_section_order(resume_data, args.sections)
    section_order = filter_section_order(
        section_order,
        include_publications=args.has_publications,
        include_volunteer=args.has_volunteer,
    )

    first_name = sanitize_filename_part(first_name_from_resume(resume_data))
    render_date = date.today().isoformat()
    file_prefix = f"{first_name}_{render_date}_{sanitize_filename_part(args.base_name)}"

    text_output = render_text_resume(resume_data, section_order)
    text_path = output_dir / f"{file_prefix}.txt"
    text_path.write_text(text_output, encoding="utf-8")

    docx_path = output_dir / f"{file_prefix}.docx"
    render_docx_resume(resume_data, section_order, docx_path)

    print(f"Wrote {text_path}")
    print(f"Wrote {docx_path}")


if __name__ == "__main__":
    main()
